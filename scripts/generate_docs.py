from pathlib import Path

from docx import Document

SOURCE_DIR = Path("docs")
OUTPUT_DIR = Path("docs/generated/docx")


def markdown_to_docx(source: Path, target: Path) -> None:
    document = Document()

    for line in source.read_text(encoding="utf-8").splitlines():
        text = line.strip()

        if not text:
            document.add_paragraph()
        elif text.startswith("# "):
            document.add_heading(text[2:], level=1)
        elif text.startswith("## "):
            document.add_heading(text[3:], level=2)
        elif text.startswith("### "):
            document.add_heading(text[4:], level=3)
        elif text.startswith("- "):
            document.add_paragraph(text[2:], style="List Bullet")
        else:
            document.add_paragraph(text)

    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(target))


def main() -> None:
    for md_file in SOURCE_DIR.rglob("*.md"):
        if "generated" in md_file.parts:
            continue

        relative_path = md_file.relative_to(SOURCE_DIR)
        target_file = OUTPUT_DIR / relative_path.with_suffix(".docx")

        markdown_to_docx(md_file, target_file)
        print(f"Generated: {target_file}")


if __name__ == "__main__":
    main()
